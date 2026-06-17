import io
from datetime import date
from types import SimpleNamespace
from typing import Any, Dict, Optional

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.models.profile import Profile


def build_tailored_profile(profile, tailored: Dict[str, Any], full_name: str, email: str):
    """
    Merge Claude's tailored CV data with the original profile ORM object,
    returning a SimpleNamespace that the existing generate_pdf / generate_docx
    functions can consume without modification.
    """
    user_ns = SimpleNamespace(full_name=full_name, email=email)

    # Map original experiences for date / location lookup
    exp_map = {(e.title or "", e.company or ""): e for e in (profile.work_experiences or [])}

    tailored_exps = []
    for te in tailored.get("experiences", []):
        key = (te.get("title", ""), te.get("company", ""))
        orig = exp_map.get(key)
        if orig:
            tailored_exps.append(SimpleNamespace(
                title=te["title"],
                company=te["company"],
                start_date=orig.start_date,
                end_date=orig.end_date,
                is_current=orig.is_current,
                location=orig.location,
                description=te.get("description", orig.description or ""),
            ))

    tailored_skills = [
        SimpleNamespace(category=sg["category"], name=name)
        for sg in tailored.get("skill_groups", [])
        for name in sg.get("names", [])
    ]

    return SimpleNamespace(
        user=user_ns,
        headline=tailored.get("headline") or profile.headline,
        phone=profile.phone,
        location=profile.location,
        linkedin_url=profile.linkedin_url,
        github_url=profile.github_url,
        summary=tailored.get("summary") or profile.summary,
        work_experiences=tailored_exps or list(profile.work_experiences or []),
        educations=list(profile.educations or []),
        skills=tailored_skills or list(profile.skills or []),
        certifications=list(profile.certifications or []),
        projects=list(profile.projects or []),
    )


def extract_text_from_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    return "\n".join(page.get_text() for page in doc)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def _fmt_date(d: Optional[date]) -> str:
    if not d:
        return ""
    return d.strftime("%b %Y")


def generate_pdf(profile: Profile) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    styles = getSampleStyleSheet()

    name_style = ParagraphStyle("Name", fontSize=20, fontName="Helvetica-Bold", spaceAfter=2)
    headline_style = ParagraphStyle("Headline", fontSize=11, fontName="Helvetica", textColor=colors.HexColor("#555555"), spaceAfter=4)
    contact_style = ParagraphStyle("Contact", fontSize=9, fontName="Helvetica", textColor=colors.HexColor("#444444"), spaceAfter=8)
    section_style = ParagraphStyle("Section", fontSize=12, fontName="Helvetica-Bold", textColor=colors.HexColor("#1a1a2e"), spaceBefore=10, spaceAfter=3, borderPadding=(0, 0, 2, 0))
    body_style = ParagraphStyle("Body", fontSize=9, fontName="Helvetica", spaceAfter=2, leading=13)
    sub_style = ParagraphStyle("Sub", fontSize=9, fontName="Helvetica-Bold", spaceAfter=1)
    meta_style = ParagraphStyle("Meta", fontSize=8, fontName="Helvetica-Oblique", textColor=colors.HexColor("#666666"), spaceAfter=2)

    story = []

    user = profile.user
    story.append(Paragraph(user.full_name, name_style))
    if profile.headline:
        story.append(Paragraph(profile.headline, headline_style))

    contact_parts = []
    if user.email:
        contact_parts.append(user.email)
    if profile.phone:
        contact_parts.append(profile.phone)
    if profile.location:
        contact_parts.append(profile.location)
    if profile.linkedin_url:
        contact_parts.append(profile.linkedin_url)
    if profile.github_url:
        contact_parts.append(profile.github_url)
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), contact_style))

    story.append(Spacer(1, 4 * mm))

    if profile.summary:
        story.append(Paragraph("SUMMARY", section_style))
        story.append(Paragraph(profile.summary.replace("\n", "<br/>"), body_style))

    if profile.work_experiences:
        story.append(Paragraph("EXPERIENCE", section_style))
        for exp in profile.work_experiences:
            date_range = _fmt_date(exp.start_date)
            if exp.is_current:
                date_range += " – Present"
            elif exp.end_date:
                date_range += f" – {_fmt_date(exp.end_date)}"
            story.append(Paragraph(f"{exp.title} @ {exp.company}", sub_style))
            meta = []
            if exp.location:
                meta.append(exp.location)
            if date_range:
                meta.append(date_range)
            if meta:
                story.append(Paragraph(" | ".join(meta), meta_style))
            if exp.description:
                for line in exp.description.split("\n"):
                    line = line.strip()
                    if line:
                        story.append(Paragraph(f"• {line}" if not line.startswith("•") else line, body_style))
            story.append(Spacer(1, 2 * mm))

    if profile.educations:
        story.append(Paragraph("EDUCATION", section_style))
        for edu in profile.educations:
            title = edu.institution
            if edu.degree:
                title = f"{edu.degree}{' in ' + edu.field_of_study if edu.field_of_study else ''} — {edu.institution}"
            story.append(Paragraph(title, sub_style))
            meta = []
            if edu.grade:
                meta.append(edu.grade)
            date_range = _fmt_date(edu.start_date)
            if edu.end_date:
                date_range += f" – {_fmt_date(edu.end_date)}"
            if date_range.strip(" –"):
                meta.append(date_range)
            if meta:
                story.append(Paragraph(" | ".join(meta), meta_style))
            story.append(Spacer(1, 1 * mm))

    if profile.projects:
        story.append(Paragraph("PERSONAL PROJECTS", section_style))
        for proj in profile.projects:
            story.append(Paragraph(proj.name, sub_style))
            meta = []
            date_range = _fmt_date(proj.start_date)
            if proj.end_date:
                date_range += f" – {_fmt_date(proj.end_date)}"
            if date_range.strip(" –"):
                meta.append(date_range)
            if proj.technologies:
                meta.append(proj.technologies)
            if meta:
                story.append(Paragraph(" | ".join(meta), meta_style))
            if proj.description:
                for line in proj.description.split("\n"):
                    line = line.strip()
                    if line:
                        story.append(Paragraph(f"• {line}" if not line.startswith("•") else line, body_style))
            links = [link for link in [proj.url, proj.repo_url] if link]
            if links:
                story.append(Paragraph(" | ".join(links), meta_style))
            story.append(Spacer(1, 2 * mm))

    if profile.skills:
        story.append(Paragraph("SKILLS", section_style))
        by_category: dict = {}
        for skill in profile.skills:
            by_category.setdefault(skill.category, []).append(skill.name)
        for cat, names in by_category.items():
            story.append(Paragraph(f"<b>{cat}:</b> {', '.join(names)}", body_style))

    if profile.certifications:
        story.append(Paragraph("CERTIFICATIONS", section_style))
        for cert in profile.certifications:
            line = cert.name
            if cert.issuer:
                line += f" — {cert.issuer}"
            if cert.issue_date:
                line += f" ({_fmt_date(cert.issue_date)})"
            story.append(Paragraph(line, body_style))

    doc.build(story)
    return buffer.getvalue()


def generate_docx(profile: Profile) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    user = profile.user

    heading = doc.add_paragraph()
    run = heading.add_run(user.full_name)
    run.bold = True
    run.font.size = Pt(18)

    if profile.headline:
        h = doc.add_paragraph(profile.headline)
        h.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    contact_parts = [x for x in [user.email, profile.phone, profile.location, profile.linkedin_url, profile.github_url] if x]
    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts))

    def add_section(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    if profile.summary:
        add_section("SUMMARY")
        doc.add_paragraph(profile.summary)

    if profile.work_experiences:
        add_section("EXPERIENCE")
        for exp in profile.work_experiences:
            p = doc.add_paragraph()
            r = p.add_run(f"{exp.title} @ {exp.company}")
            r.bold = True
            date_range = _fmt_date(exp.start_date)
            if exp.is_current:
                date_range += " – Present"
            elif exp.end_date:
                date_range += f" – {_fmt_date(exp.end_date)}"
            meta = [x for x in [exp.location, date_range] if x]
            if meta:
                m = doc.add_paragraph(" | ".join(meta))
                m.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                m.runs[0].font.italic = True
            if exp.description:
                doc.add_paragraph(exp.description)

    if profile.educations:
        add_section("EDUCATION")
        for edu in profile.educations:
            title = edu.institution
            if edu.degree:
                title = f"{edu.degree}{' in ' + edu.field_of_study if edu.field_of_study else ''} — {edu.institution}"
            p = doc.add_paragraph()
            p.add_run(title).bold = True
            meta = []
            if edu.grade:
                meta.append(edu.grade)
            date_range = _fmt_date(edu.start_date)
            if edu.end_date:
                date_range += f" – {_fmt_date(edu.end_date)}"
            if date_range.strip(" –"):
                meta.append(date_range)
            if meta:
                doc.add_paragraph(" | ".join(meta))

    if profile.projects:
        add_section("PERSONAL PROJECTS")
        for proj in profile.projects:
            p = doc.add_paragraph()
            p.add_run(proj.name).bold = True
            date_range = _fmt_date(proj.start_date)
            if proj.end_date:
                date_range += f" – {_fmt_date(proj.end_date)}"
            meta = [x for x in [date_range.strip(" –") or None, proj.technologies] if x]
            if meta:
                m = doc.add_paragraph(" | ".join(meta))
                m.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                m.runs[0].font.italic = True
            if proj.description:
                doc.add_paragraph(proj.description)
            links = [link for link in [proj.url, proj.repo_url] if link]
            if links:
                doc.add_paragraph(" | ".join(links))

    if profile.skills:
        add_section("SKILLS")
        by_category: dict = {}
        for skill in profile.skills:
            by_category.setdefault(skill.category, []).append(skill.name)
        for cat, names in by_category.items():
            p = doc.add_paragraph()
            p.add_run(f"{cat}: ").bold = True
            p.add_run(", ".join(names))

    if profile.certifications:
        add_section("CERTIFICATIONS")
        for cert in profile.certifications:
            line = cert.name
            if cert.issuer:
                line += f" — {cert.issuer}"
            if cert.issue_date:
                line += f" ({_fmt_date(cert.issue_date)})"
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
